"""Generate rich sleep analysis reports as HTML or DOCX."""
from datetime import datetime
import base64
import os
from io import BytesIO

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.font_manager as fm
import pandas as pd
import numpy as np

# --- Chinese font setup ---
import platform as _platform

_cjk_font_path = None

# 1) Try well-known paths first (works even if matplotlib cache is stale)
_well_known = [
    "C:/Windows/Fonts/simhei.ttf",
    "C:/Windows/Fonts/msyh.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    "/usr/share/fonts/truetype/arphic/uming.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
]
for _p in _well_known:
    if os.path.isfile(_p):
        _cjk_font_path = _p
        break

# 2) Fallback: scan with matplotlib
if _cjk_font_path is None:
    try:
        for _f in fm.findSystemFonts():
            _fl = _f.lower()
            if any(_k in _fl for _k in ["simhei", "msyh", "yahei", "wqy", "wenquan",
                                          "noto sans cjk", "notosanscjk", "uming", "gkai"]):
                _cjk_font_path = _f
                break
    except Exception:
        pass

if _cjk_font_path is not None:
    try:
        fm.fontManager.addfont(_cjk_font_path)
        _n = fm.FontProperties(fname=_cjk_font_path).get_name()
        plt.rcParams["font.family"] = "sans-serif"
        plt.rcParams["font.sans-serif"] = [_n] + [
            f for f in plt.rcParams.get("font.sans-serif", []) if f != _n
        ]
    except Exception:
        pass

plt.rcParams["axes.unicode_minus"] = False

# Warm palette for sleep stages
STAGE_COLORS = {0: "#E8904C", 1: "#5DAF8B", 2: "#8B7EC8"}  # Wake=warm amber, NREM=sage, REM=soft purple
STAGE_NAMES = {0: "清醒", 1: "深睡眠", 2: "做梦期"}


def plot_hypnogram(predictions: pd.DataFrame, figsize=(12, 1.8)) -> plt.Figure:
    """Plot sleep stage hypnogram — clean horizontal bar matching SPA canvas style."""
    labels = predictions["predicted_label"].values
    n_epochs = len(labels)
    total_hours = n_epochs * 30 / 3600
    total_minutes = int(total_hours * 60)

    fig, ax = plt.subplots(figsize=figsize, facecolor="#FEFAF5")
    ax.set_facecolor("#FEFAF5")

    # Draw each 30 s epoch as a thin colored vertical span
    for i in range(n_epochs):
        ax.axvspan(i, i + 1, facecolor=STAGE_COLORS.get(labels[i], "#999999"),
                   alpha=0.88, linewidth=0)

    ax.set_xlim(0, n_epochs)
    ax.set_ylim(0, 1)
    ax.set_yticks([])

    # X-axis: time labels (relative hours like the SPA canvas fallback)
    if total_hours >= 2:
        hour_ticks = list(range(0, int(total_hours) + 1))
        tick_positions = [h * 3600 / 30 for h in hour_ticks]
        tick_labels = [f"{h}h" for h in hour_ticks]
    else:
        tick_positions = [0, n_epochs]
        tick_labels = ["0", f"{total_minutes}m"]

    ax.set_xticks(tick_positions)
    ax.set_xticklabels(tick_labels, fontsize=12, color="#6D5C4F")
    ax.tick_params(bottom=True, left=False, colors="#6D5C4F", labelsize=11)

    ax.set_title("整晚睡眠一览", color="#3E2E22", fontweight="bold", fontsize=18, pad=10)

    # Clean frame — only a subtle bottom line
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.spines["bottom"].set_visible(True)
    ax.spines["bottom"].set_color("#D7C4B0")

    # Legend below the bar (matching SPA canvas position)
    legend_patches = [
        mpatches.Patch(color=STAGE_COLORS[0], label="清醒"),
        mpatches.Patch(color=STAGE_COLORS[1], label="深睡眠"),
        mpatches.Patch(color=STAGE_COLORS[2], label="做梦期"),
    ]
    ax.legend(handles=legend_patches, loc="lower left", fontsize=11,
              ncol=3, frameon=False, bbox_to_anchor=(0, -0.28))

    # Add duration label on the right
    hrs = int(total_hours)
    mins = total_minutes % 60
    dur_text = f"{hrs}h{mins}m" if mins > 0 else f"{hrs}h"
    ax.text(1.005, 0.5, f"共 {dur_text}", transform=ax.transAxes,
            fontsize=10, color="#A38B78", va="center")

    plt.tight_layout()
    return fig


def plot_stage_distribution(predictions: pd.DataFrame, figsize=(5, 4)) -> plt.Figure:
    """Plot sleep stage distribution as a donut chart."""
    labels = predictions["predicted_label"].values
    counts = pd.Series(labels).value_counts().sort_index()
    total = len(labels)

    fig, ax = plt.subplots(figsize=figsize, facecolor="#FEFAF5")
    stage_labels = [f"{STAGE_NAMES.get(i, '?')}" for i in [0, 1, 2]]
    colors = [STAGE_COLORS[i] for i in [0, 1, 2]]
    sizes = [counts.get(i, 0) for i in [0, 1, 2]]

    wedges, texts, autotexts = ax.pie(
        sizes, labels=stage_labels, colors=colors, autopct="%1.1f%%",
        startangle=90, pctdistance=0.75,
        wedgeprops=dict(width=0.4, edgecolor='white', linewidth=2),
    )
    for at in autotexts:
        at.set_fontsize(14)
        at.set_fontweight("bold")
        at.set_color("#3E2E22")
    for t in texts:
        t.set_fontsize(14)
        t.set_color("#6D5C4F")

    ax.set_title("睡眠阶段占比", color="#3E2E22", fontweight="bold", fontsize=16)
    plt.tight_layout()
    return fig


def plot_metric_summary(metrics: dict, ref_comparison: dict = None, figsize=(8, 4)) -> plt.Figure:
    """Plot a horizontal bar chart comparing user metrics to reference ranges."""
    if ref_comparison is None:
        from src.sleep_scoring import generate_reference_comparison
        ref_comparison = generate_reference_comparison(metrics)

    friendly_keys = {
        "睡眠效率 SE (%)": "睡眠质量",
        "总睡眠时长 TST (分钟)": "睡眠时长",
        "入睡后清醒 WASO (分钟)": "中途清醒",
        "入睡潜伏期 (分钟)": "入睡速度",
        "REM 占比 (%)": "做梦比例",
    }
    plot_metrics = ["睡眠效率 SE (%)", "总睡眠时长 TST (分钟)", "入睡后清醒 WASO (分钟)",
                    "入睡潜伏期 (分钟)", "REM 占比 (%)"]
    available = [k for k in plot_metrics if k in ref_comparison]

    fig, axes = plt.subplots(len(available), 1, figsize=figsize, facecolor="#FEFAF5")
    if len(available) == 1:
        axes = [axes]

    for idx, key in enumerate(available):
        ax = axes[idx]
        rc = ref_comparison[key]
        val = rc["your_value"]
        if isinstance(val, str):
            val = 0
        val = float(val)

        status_colors = {"normal": "#5DAF8B", "borderline": "#E8904C", "concerning": "#D35D47", "unknown": "#AAA"}
        color = status_colors.get(rc["status"], "#AAA")
        status_labels = {"normal": "正常", "borderline": "偏高/偏低", "concerning": "需关注", "unknown": "未知"}
        label = status_labels.get(rc["status"], "")
        friendly_name = friendly_keys.get(key, key)

        ax.barh(0, val, color=color, height=0.5, alpha=0.8)
        ax.set_xlim(0, val * 1.5 if val > 0 else 100)
        ax.set_yticks([])
        ax.set_title(f"{friendly_name}: {val}  [{label}]", fontsize=13, color="#3E2E22", loc="left")
        ax.axvline(x=float(val), color=color, linewidth=2)

        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_visible(False)
        ax.spines["bottom"].set_color("#D7C4B0")
        ax.tick_params(colors="#6D5C4F", labelsize=11)

    fig.suptitle("核心指标", fontsize=16, color="#3E2E22", fontweight="bold", y=1.02)
    plt.tight_layout()
    return fig


def fig_to_base64(fig: plt.Figure) -> str:
    """Convert matplotlib figure to base64 PNG string."""
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    buf.seek(0)
    plt.close(fig)
    return base64.b64encode(buf.read()).decode()


def generate_html_report(
    predictions: pd.DataFrame,
    metrics: dict,
    sleep_score: dict = None,
    recommendations: list = None,
    ref_comparison: dict = None,
    shap_importance: pd.DataFrame = None,
    explanation_text: str = "",
    upload_filename: str = "",
    format_metadata: dict = None,
) -> str:
    """Generate a rich HTML sleep report with scoring and recommendations.

    Args:
        predictions: DataFrame with predicted_label column
        metrics: dict from compute_sleep_metrics()
        sleep_score: dict from compute_sleep_score()
        recommendations: list from generate_recommendations()
        ref_comparison: dict from generate_reference_comparison()
        shap_importance: optional SHAP importance DataFrame
        explanation_text: natural language SHAP explanation
        upload_filename: original filename
        format_metadata: dict with source_format, conversion_notes, features_real, features_synthesized
    """
    # Generate charts
    hypo_fig = plot_hypnogram(predictions)
    hypo_b64 = fig_to_base64(hypo_fig)

    dist_fig = plot_stage_distribution(predictions)
    dist_b64 = fig_to_base64(dist_fig)

    # Compute scoring if not provided
    if sleep_score is None:
        from src.sleep_scoring import compute_sleep_score, generate_recommendations, generate_reference_comparison, generate_summary_text
        sleep_score = compute_sleep_score(metrics)
    if recommendations is None:
        from src.sleep_scoring import generate_recommendations
        recommendations = generate_recommendations(metrics, sleep_score.get("subscores"))
    if ref_comparison is None:
        from src.sleep_scoring import generate_reference_comparison
        ref_comparison = generate_reference_comparison(metrics)

    from src.sleep_scoring import generate_summary_text
    summary = generate_summary_text(sleep_score, recommendations)

    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    total = sleep_score.get("total_score", 0)
    grade = sleep_score.get("grade", "?")
    grade_label = sleep_score.get("grade_label", "?")

    # Grade badge color
    grade_colors = {
        "A+": "#2E7D32", "A": "#388E3C", "B+": "#558B2F", "B": "#689F38",
        "C+": "#F9A825", "C": "#F57F17", "D": "#E65100", "F": "#BF360C",
    }
    badge_color = grade_colors.get(grade, "#999")

    # Score gauge: simple SVG circular gauge
    gauge_svg = _generate_score_gauge_svg(total, badge_color)

    # Metrics table — with friendly names
    friendly_names = {
        "总睡眠时长 TST (分钟)": "睡了多久",
        "睡眠效率 SE (%)": "睡眠质量",
        "入睡后清醒 WASO (分钟)": "中途清醒",
        "入睡潜伏期 (分钟)": "多久入睡",
        "REM 潜伏期 (分钟)": "做梦潜伏期",
        "NREM 占比 (%)": "深睡眠比例",
        "REM 占比 (%)": "做梦比例",
        "阶段转换次数": "阶段转换",
        "睡眠周期数": "睡眠周期",
        "总记录时长 (分钟)": "总记录时长",
    }
    metrics_rows = ""
    status_icons = {"normal": "&#10003;", "borderline": "&#9888;", "concerning": "&#10007;", "unknown": "&#9679;"}
    status_colors_css = {"normal": "#5DAF8B", "borderline": "#E8904C", "concerning": "#D35D47", "unknown": "#AAA"}
    status_labels_zh = {"normal": "正常", "borderline": "偏高/偏低", "concerning": "需关注", "unknown": "未知"}

    for key, rc in ref_comparison.items():
        status = rc["status"]
        icon = status_icons.get(status, "")
        color = status_colors_css.get(status, "#AAA")
        label = status_labels_zh.get(status, "")
        friendly_key = friendly_names.get(key, key)
        metrics_rows += f"""
        <tr>
            <td style="padding:10px 14px;border-bottom:1px solid #EDE0D4;">
                <strong>{friendly_key}</strong>
                <div style="font-size:0.78rem;color:#A38B78;">{rc['interpretation']}</div>
            </td>
            <td style="padding:10px 14px;border-bottom:1px solid #EDE0D4;text-align:center;font-weight:700;font-size:1.1rem;">
                {rc['your_value']}
            </td>
            <td style="padding:10px 14px;border-bottom:1px solid #EDE0D4;text-align:center;color:#A38B78;font-size:0.85rem;">
                {rc['reference_range']}
            </td>
            <td style="padding:10px 14px;border-bottom:1px solid #EDE0D4;text-align:center;">
                <span style="color:{color};font-weight:600;">{icon} {label}</span>
            </td>
        </tr>"""

    # Recommendations section — conversational tone
    icon_map = {
        "efficiency": "&#128164;", "fragmentation": "&#128564;", "latency": "&#9200;",
        "rem": "&#127752;", "general": "&#128161;", "duration": "&#9200;",
        "insomnia_pattern": "&#128564;", "fragmentation_combo": "&#128564;",
        "sleep_restriction": "&#9200;", "long_bed": "&#128164;",
        "recovery_sleep": "&#127752;", "alcohol_effect": "&#127863;",
    }
    recs_html = ""
    for i, rec in enumerate(recommendations):
        sev_colors = {"critical": "#D35D47", "warning": "#E8904C", "info": "#5DAF8B"}
        cat = rec.get("category", "general")
        icon = icon_map.get(cat, icon_map["general"])
        advice = rec["advice"]
        # Conversational rewrite based on category
        if cat == "efficiency":
            se_val = float(metrics.get("睡眠效率 SE (%)", 85))
            advice = f"你的睡眠效率是 {se_val:.0f}%，躺在床上的时间有 {100 - int(se_val)}% 没有真正睡着。建议只在困了才上床，醒了就起来，别在床上刷手机。"
        elif cat == "fragmentation":
            waso_val = float(metrics.get("入睡后清醒 WASO (分钟)", 30))
            advice = f"你昨晚中途醒了 {waso_val:.0f} 分钟，睡眠被打断了。睡前少喝水、保持房间安静黑暗，可能会有帮助。"
        elif cat == "latency":
            lat_val = float(metrics.get("入睡潜伏期 (分钟)", 15))
            advice = f"你花了 {lat_val:.0f} 分钟才入睡，有点慢。睡前一小时放下手机，试试深呼吸或泡个热水脚。"
        elif cat == "rem":
            rem_val = float(metrics.get("REM 占比 (%)", 22))
            if rem_val > 28:
                advice = f"你的做梦期占比 {rem_val:.0f}%，高于正常的 20-25%。这在医学上叫「REM 反弹」——身体在补偿之前被压制的做梦时间。常见原因：① 近期睡眠不足，补觉时优先补回 REM；② 戒酒或停用某些药物（如抗抑郁药、安眠药）；③ 近期经历了较大压力，大脑通过增加 REM 来调节情绪。REM 反弹是身体的自我修复机制，短期出现属于正常。如果持续偏高，建议保持规律作息、减少饮酒，必要时咨询医生。"
            else:
                advice = f"你的做梦期只占 {rem_val:.0f}%，低于正常的 20-25%。做梦期是大脑整理记忆和调节情绪的关键阶段——少了它，第二天容易健忘、情绪波动。常见原因：饮酒（哪怕一杯也会压制做梦）、压力大、作息不规律。试着睡前放松、减少饮酒，做梦时间会慢慢恢复。"
        recs_html += f"""
        <div style="background:#FEFAF5;border:1px solid #EDE0D4;border-radius:12px;padding:16px;margin-bottom:12px;border-left:4px solid {sev_colors.get(rec['severity'], '#999')};">
            <div style="display:flex;align-items:flex-start;gap:12px;">
                <span style="font-size:1.6rem;line-height:1;">{icon}</span>
                <div style="flex:1;">
                    <div style="font-weight:600;color:#3E2E22;margin-bottom:4px;">{rec['issue']}</div>
                    <p style="color:#6D5C4F;line-height:1.7;margin:0;">{advice}</p>
                </div>
            </div>
        </div>"""

    # Data quality notes
    quality_html = ""
    if format_metadata:
        notes = format_metadata.get("conversion_notes", [])
        real = format_metadata.get("features_real", [])
        synth = format_metadata.get("features_synthesized", [])
        quality_html = "<ul style='color:#6D5C4F;line-height:1.8;'>"
        for note in notes:
            quality_html += f"<li>{note}</li>"
        if real:
            quality_html += f"<li>真实传感器数据: {', '.join(real)}</li>"
        if synth:
            quality_html += f"<li>算法估算数据: {', '.join(synth)} (Apple Watch 不导出原始加速度数据)</li>"
        quality_html += "</ul>"

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>睡眠分析报告</title>
<link href="https://fonts.googleapis.com/css2?family=Sora:wght@500;600;700&family=Outfit:wght@400;500;600&display=swap" rel="stylesheet">
<style>
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{
    font-family: 'Outfit', 'Microsoft YaHei', 'PingFang SC', sans-serif;
    background: oklch(0.97 0.012 85);
    color: oklch(0.22 0.02 75);
    line-height: 1.7;
  }}
  .report-container {{
    max-width: 780px;
    margin: 0 auto;
    padding: 48px 24px 80px;
  }}
  .cover {{
    text-align: center;
    padding: 48px 24px;
    background: #FEFAF5;
    border-radius: 20px;
    margin-bottom: 32px;
    box-shadow: 0 2px 16px rgba(0,0,0,0.06);
  }}
  .cover h1 {{
    font-family: 'Sora', sans-serif;
    font-size: 2rem;
    font-weight: 700;
    color: oklch(0.22 0.02 75);
    margin-bottom: 8px;
  }}
  .cover .meta {{
    color: oklch(0.50 0.02 75);
    font-size: 0.85rem;
  }}
  .score-section {{
    text-align: center;
    padding: 40px 24px;
    background: #FEFAF5;
    border-radius: 20px;
    margin-bottom: 24px;
    box-shadow: 0 2px 16px rgba(0,0,0,0.06);
  }}
  .score-gauge {{
    display: flex;
    justify-content: center;
    margin-bottom: 16px;
  }}
  .grade-badge {{
    display: inline-block;
    background: {badge_color};
    color: #fff;
    padding: 6px 20px;
    border-radius: 100px;
    font-family: 'Sora', sans-serif;
    font-size: 1.5rem;
    font-weight: 700;
    margin: 8px 0;
  }}
  .headline {{
    font-size: 1.3rem;
    font-weight: 600;
    color: oklch(0.22 0.02 75);
    margin-bottom: 4px;
  }}
  .section {{
    background: #FEFAF5;
    border-radius: 16px;
    padding: 28px;
    margin-bottom: 20px;
    box-shadow: 0 2px 12px rgba(0,0,0,0.05);
  }}
  .section h3 {{
    font-family: 'Sora', sans-serif;
    font-size: 1.1rem;
    font-weight: 600;
    color: oklch(0.22 0.02 75);
    margin-bottom: 16px;
    padding-bottom: 10px;
    border-bottom: 2px solid #EDE0D4;
  }}
  .metrics-table {{
    width: 100%;
    border-collapse: collapse;
  }}
  .metrics-table th {{
    padding: 10px 14px;
    text-align: center;
    font-size: 0.8rem;
    color: #A38B78;
    font-weight: 600;
    border-bottom: 2px solid #EDE0D4;
  }}
  .metrics-table tr:hover {{
    background: rgba(237, 224, 212, 0.3);
  }}
  img {{
    max-width: 100%;
    border-radius: 12px;
  }}
  .subscore-bar {{
    display: flex;
    align-items: center;
    gap: 10px;
    margin-bottom: 8px;
  }}
  .subscore-label {{
    width: 140px;
    font-size: 0.82rem;
    color: #6D5C4F;
    text-align: right;
  }}
  .subscore-track {{
    flex: 1;
    height: 8px;
    background: #EDE0D4;
    border-radius: 4px;
    overflow: hidden;
  }}
  .subscore-fill {{
    height: 100%;
    border-radius: 4px;
    background: oklch(0.58 0.17 35);
  }}
  .subscore-val {{
    width: 60px;
    font-size: 0.78rem;
    color: #A38B78;
    font-weight: 600;
  }}
  .plain-summary {{
    background: linear-gradient(135deg, #FFF8F0, #FEFAF5);
    border: 1px solid #EDE0D4;
    border-radius: 14px;
    padding: 20px 24px;
    font-size: 1.05rem;
    line-height: 1.9;
    color: #5D4037;
  }}
  .disclaimer {{
    background: #FFF8E1;
    border: 2px solid #FFB74D;
    border-radius: 12px;
    padding: 20px;
    margin-top: 24px;
    color: #5D4037;
    font-size: 0.85rem;
    line-height: 1.7;
  }}
  .disclaimer strong {{
    color: #E65100;
  }}
  .footer {{
    text-align: center;
    color: #A38B78;
    font-size: 0.75rem;
    margin-top: 32px;
    padding-top: 20px;
    border-top: 1px solid #EDE0D4;
  }}
</style>
</head>
<body>
<div class="report-container">

  <!-- Cover -->
  <div class="cover">
    <h1>我的睡眠报告</h1>
    <div class="meta">
      生成时间: {now_str} &nbsp;|&nbsp; 数据来源: {upload_filename or "上传数据"}
    </div>
  </div>

  <!-- Sleep Score -->
  <div class="score-section">
    <div class="headline">{summary.get('headline', '')}</div>
    <div class="score-gauge">{gauge_svg}</div>
    <div class="grade-badge">{grade} · {grade_label}</div>
    <p style="color:#A38B78;font-size:0.9rem;margin-top:4px;">分数越高，睡眠越好</p>
    <p style="color:#6D5C4F;margin-top:12px;max-width:500px;margin-left:auto;margin-right:auto;">{summary.get('plain_summary', '')}</p>
    <!-- Subscores -->
    <div style="margin-top:24px;max-width:420px;margin-left:auto;margin-right:auto;">
      {_generate_subscore_bars_html(sleep_score.get('subscores', {}))}
    </div>
  </div>

  <!-- Hypnogram -->
  <div class="section">
    <h3>整晚睡眠一览</h3>
    <img src="data:image/png;base64,{hypo_b64}" alt="睡眠分期图">
  </div>

  <!-- Stage Distribution -->
  <div class="section">
    <h3>睡眠阶段分布</h3>
    <img src="data:image/png;base64,{dist_b64}" alt="阶段分布" style="max-width:360px;display:block;margin:0 auto;">
  </div>

  <!-- Metrics Dashboard -->
  <div class="section">
    <h3>核心指标</h3>
    <table class="metrics-table">
      <thead><tr>
        <th>指标</th>
        <th>你的数值</th>
        <th>正常范围</th>
        <th>状态</th>
      </tr></thead>
      <tbody>{metrics_rows}</tbody>
    </table>
    <div style="background:#F5F0FF;border:1px solid #D4C8ED;border-radius:10px;padding:14px 18px;margin-top:14px;">
      <div style="font-weight:600;color:#5B4A9E;margin-bottom:6px;">&#127752; 关于"做梦时间"（REM）</div>
      <div style="color:#5D4037;font-size:0.92rem;line-height:1.8;">
        做梦主要发生在 REM（快速眼动）睡眠阶段。REM 对<strong>记忆巩固</strong>和<strong>情绪调节</strong>非常重要——就像大脑在夜间"整理文件"和"清理情绪垃圾"。<br>
        正常范围是占总睡眠的 <strong>20-25%</strong>。偏少（&lt;18%）常见于饮酒、压力大或作息紊乱；偏多（&gt;28%）在医学上称为「REM 反弹」——身体在补偿之前被压制的做梦时间，常见于近期睡眠不足、戒酒或压力较大时。短期属于正常的自我修复机制，持续偏高建议咨询医生。
      </div>
    </div>
  </div>

  <!-- Recommendations -->
  <div class="section">
    <h3>睡眠小贴士</h3>
    {recs_html}
  </div>

  <!-- Key Insights -->
  <div class="section">
    <h3>关键发现</h3>
    <div class="plain-summary">
      {_generate_plain_language_extra_html(metrics)}
    </div>
  </div>
"""

    # SHAP Explanation
    if explanation_text:
        html += f"""
  <div class="section">
    <h3>模型如何做出的判断</h3>
    <div style="color:#6D5C4F;line-height:1.9;">{explanation_text}</div>
  </div>"""

    # Data Quality
    if quality_html:
        html += f"""
  <div class="section">
    <h3>数据质量说明</h3>
    {quality_html}
  </div>"""

    # References
    html += """
  <div class="section">
    <h3>参考文献</h3>
    <ul style="color:#6D5C4F;font-size:0.85rem;line-height:2;">
      <li>American Academy of Sleep Medicine (AASM) Clinical Practice Guidelines</li>
      <li>Sleep Foundation — Sleep Hygiene Recommendations</li>
      <li>National Institutes of Health (NIH) — Brain Basics: Understanding Sleep</li>
      <li>CDC — Sleep and Sleep Disorders</li>
      <li>Sleep-Accel Dataset (PhysioNet) — Wearable-based sleep staging</li>
    </ul>
  </div>
"""

    # Disclaimer
    html += f"""
  <div class="disclaimer">
    <strong>温馨提示</strong><br>
    本报告基于你的可穿戴设备数据，通过算法估算你的睡眠情况。它不能代替医院的专业睡眠检测（多导睡眠图）。<br>
    如果你有长期睡眠问题（如严重失眠、打鼾、白天嗜睡），建议咨询睡眠专科医生。
  </div>

  <div class="footer">
    生成于 {now_str} &nbsp;|&nbsp; 仅供健康管理参考，不构成医疗建议
  </div>

</div>
</body>
</html>"""
    return html


def _generate_score_gauge_svg(score: int, color: str) -> str:
    """Generate an SVG circular gauge for the sleep score."""
    r = 60
    circumference = 2 * 3.14159 * r
    offset = circumference * (1 - score / 100)

    return f"""
    <svg width="160" height="100" viewBox="0 0 160 100">
      <defs>
        <linearGradient id="scoreGrad" x1="0%" y1="0%" x2="100%" y2="0%">
          <stop offset="0%" style="stop-color:#D35D47;stop-opacity:1" />
          <stop offset="50%" style="stop-color:#E8904C;stop-opacity:1" />
          <stop offset="100%" style="stop-color:#5DAF8B;stop-opacity:1" />
        </linearGradient>
      </defs>
      <circle cx="80" cy="80" r="{r}" fill="none" stroke="#EDE0D4" stroke-width="12"
              stroke-dasharray="{circumference}" stroke-dashoffset="0"
              transform="rotate(180 80 80)" />
      <circle cx="80" cy="80" r="{r}" fill="none" stroke="url(#scoreGrad)" stroke-width="12"
              stroke-dasharray="{circumference}" stroke-dashoffset="{offset}"
              stroke-linecap="round" transform="rotate(180 80 80)" />
      <text x="80" y="75" text-anchor="middle" font-family="Sora,sans-serif"
            font-size="36" font-weight="700" fill="{color}">{score}</text>
      <text x="80" y="92" text-anchor="middle" font-family="Outfit,sans-serif"
            font-size="11" fill="#A38B78">/ 100</text>
    </svg>"""


def _generate_subscore_bars_html(subscores: dict) -> str:
    """Generate HTML for subscore progress bars."""
    labels = {
        "sleep_efficiency": "睡眠质量",
        "total_sleep_time": "睡眠时长",
        "waso": "中途清醒",
        "sleep_latency": "入睡速度",
        "rem_proportion": "做梦比例",
    }
    icons = {
        "sleep_efficiency": "&#127775;",
        "total_sleep_time": "&#128164;",
        "waso": "&#128564;",
        "sleep_latency": "&#9200;",
        "rem_proportion": "&#127752;",
    }
    html = ""
    for key, data in subscores.items():
        label = labels.get(key, key)
        icon = icons.get(key, "&#9679;")
        score = data.get("score", 0)
        max_s = data.get("max", 1)
        pct = score / max_s * 100 if max_s > 0 else 0
        if pct >= 80:
            status_text, status_color = "优秀", "#5DAF8B"
        elif pct >= 60:
            status_text, status_color = "良好", "#5DAF8B"
        elif pct >= 40:
            status_text, status_color = "一般", "#E8904C"
        else:
            status_text, status_color = "待改善", "#D35D47"
        html += f"""
        <div class="subscore-bar">
          <span style="font-size:1.1rem;">{icon}</span>
          <span class="subscore-label">{label}</span>
          <div class="subscore-track"><div class="subscore-fill" style="width:{pct}%;background:{status_color}"></div></div>
          <span class="subscore-val" style="color:{status_color};font-weight:600;">{status_text}</span>
        </div>"""
    return html


def _generate_plain_language_extra_html(metrics: dict) -> str:
    """Generate extra plain-language explanations for key metrics."""
    parts = []

    tst = metrics.get("总睡眠时长 TST (分钟)", 0)
    if isinstance(tst, (int, float)) and tst > 0:
        hrs = int(tst // 60)
        mins = int(tst % 60)
        parts.append(f"你昨晚睡了 <strong>{hrs} 小时 {mins} 分钟</strong>。")

    se = metrics.get("睡眠效率 SE (%)", 0)
    if isinstance(se, (int, float)) and se > 0:
        if se >= 90:
            parts.append("睡眠效率很高，躺在床上的时间大部分都用在了实际睡眠上。")
        elif se >= 80:
            parts.append("睡眠效率还不错，但还有提升空间。")
        else:
            parts.append("躺在床上的时间有一部分没睡着，可以试试只在困了才上床。")

    rem = metrics.get("REM 占比 (%)", 0)
    if isinstance(rem, (int, float)):
        parts.append(f"做梦期占 {rem:.0f}%，这个阶段对记忆和情绪调节很重要。")

    transitions = metrics.get("阶段转换次数", 0)
    if isinstance(transitions, (int, float)) and transitions > 60:
        parts.append("您的睡眠阶段转换次数偏多，说明夜间不太安稳。")

    return "<p>" + "</p><p>".join(parts) + "</p>" if parts else ""


def generate_docx_report(
    predictions: pd.DataFrame,
    metrics: dict,
    sleep_score: dict = None,
    recommendations: list = None,
    ref_comparison: dict = None,
    shap_importance: pd.DataFrame = None,
    explanation_text: str = "",
    upload_filename: str = "",
    format_metadata: dict = None,
) -> BytesIO:
    """Generate a DOCX sleep report with scoring and recommendations."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return None

    # Compute scoring if needed
    if sleep_score is None:
        from src.sleep_scoring import compute_sleep_score, generate_recommendations, generate_reference_comparison
        sleep_score = compute_sleep_score(metrics)
    if recommendations is None:
        from src.sleep_scoring import generate_recommendations
        recommendations = generate_recommendations(metrics, sleep_score.get("subscores"))
    if ref_comparison is None:
        from src.sleep_scoring import generate_reference_comparison
        ref_comparison = generate_reference_comparison(metrics)

    from src.sleep_scoring import generate_summary_text
    summary = generate_summary_text(sleep_score, recommendations)

    doc = Document()

    # Style setup
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.5

    # Title
    title = doc.add_heading("睡眠分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"生成时间: {now_str} | 数据来源: {upload_filename or '上传数据'}\n"
        f"模型版本: HGB + Context (±2) | 基于 Sleep-Accel 公开数据集训练"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x9E, 0x8B, 0x78)

    doc.add_paragraph()

    # Sleep Score
    doc.add_heading("睡眠评分", 1)
    total = sleep_score.get("total_score", 0)
    grade = sleep_score.get("grade", "?")
    grade_label = sleep_score.get("grade_label", "?")
    score_p = doc.add_paragraph()
    score_run = score_p.add_run(f"{total} 分  |  {grade} · {grade_label}")
    score_run.font.size = Pt(18)
    score_run.bold = True
    doc.add_paragraph(summary.get("plain_summary", ""))

    # Hypnogram
    doc.add_heading("整夜睡眠分期图", 1)
    hypo_fig = plot_hypnogram(predictions, figsize=(8, 3))
    hypo_buf = BytesIO()
    hypo_fig.savefig(hypo_buf, format="png", dpi=150, bbox_inches="tight", facecolor=hypo_fig.get_facecolor())
    hypo_buf.seek(0)
    plt.close(hypo_fig)
    doc.add_picture(hypo_buf, width=Inches(6))

    # Stage Distribution
    doc.add_heading("睡眠阶段分布", 1)
    dist_fig = plot_stage_distribution(predictions)
    dist_buf = BytesIO()
    dist_fig.savefig(dist_buf, format="png", dpi=150, bbox_inches="tight", facecolor=dist_fig.get_facecolor())
    dist_buf.seek(0)
    plt.close(dist_fig)
    doc.add_picture(dist_buf, width=Inches(3.5))

    # Metrics table
    doc.add_heading("指标仪表盘", 1)
    table = doc.add_table(rows=len(ref_comparison) + 1, cols=4)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["指标", "您的数值", "参考范围", "状态"]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    for i, (key, rc) in enumerate(ref_comparison.items(), 1):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = str(rc["your_value"])
        table.cell(i, 2).text = rc["reference_range"]
        status_text = {"normal": "正常", "borderline": "临界", "concerning": "需关注"}.get(rc["status"], "未知")
        table.cell(i, 3).text = f"{status_text} — {rc['interpretation']}"

    # Recommendations
    doc.add_heading("个性化改善建议", 1)
    for i, rec in enumerate(recommendations):
        p = doc.add_paragraph()
        run = p.add_run(f"{i+1}. {rec['issue']}")
        run.bold = True
        run.font.size = Pt(10)
        doc.add_paragraph(rec["advice"])
        ref_p = doc.add_paragraph()
        ref_run = ref_p.add_run(f"参考: {rec.get('reference', '')}")
        ref_run.font.size = Pt(8)
        ref_run.font.color.rgb = RGBColor(0xA3, 0x8B, 0x78)

    # Data quality
    if format_metadata:
        doc.add_heading("数据质量说明", 1)
        for note in format_metadata.get("conversion_notes", []):
            doc.add_paragraph(note, style="List Bullet")

    # Explanation
    if explanation_text:
        doc.add_heading("模型如何做出判断", 1)
        clean_text = explanation_text.replace("**", "").replace("*", "")
        doc.add_paragraph(clean_text)

    # Disclaimer
    doc.add_heading("重要说明", 1)
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "本报告基于可穿戴设备信号进行算法估计，不等同于临床多导睡眠图（PSG）诊断。\n"
        "当前模型为整晚数据上传后的离线睡眠分期分析，使用了未来 epoch 的上下文信息（±2 窗口），不属于实时分析。\n"
        "模型在 Sleep-Accel 公开数据集（31 名受试者）上训练，未进行外部验证。本结果仅供健康管理参考，不构成医疗诊断。\n"
        "如有睡眠问题困扰，请咨询睡眠专科医生。"
    )
    disclaimer_run.font.size = Pt(9)
    disclaimer_run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Evidence-focused report generators
# ---------------------------------------------------------------------------
# These definitions intentionally override the earlier richer report functions.
# The generated reports now focus on model outputs: staging predictions,
# stage distribution, and directly computed sleep metrics. Custom quality
# scores, grades, and recommendation rules are not included.


def _stage_summary(predictions: pd.DataFrame) -> list:
    labels = predictions["predicted_stage"].astype(str)
    total = len(labels)
    rows = []
    for stage in ["Wake", "NREM", "REM"]:
        count = int((labels == stage).sum())
        minutes = round(count * 30 / 60, 1)
        pct = round(count / total * 100, 1) if total else 0
        rows.append({"stage": stage, "epochs": count, "minutes": minutes, "pct": pct})
    return rows


def _html_escape(value) -> str:
    import html
    return html.escape(str(value))


def generate_html_report(
    predictions: pd.DataFrame,
    metrics: dict,
    sleep_score: dict = None,
    recommendations: list = None,
    ref_comparison: dict = None,
    shap_importance: pd.DataFrame = None,
    explanation_text: str = "",
    upload_filename: str = "",
    format_metadata: dict = None,
) -> str:
    """Generate an evidence-focused HTML report without scores or advice."""
    hypo_b64 = fig_to_base64(plot_hypnogram(predictions))
    dist_b64 = fig_to_base64(plot_stage_distribution(predictions))
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    stage_rows = "".join(
        f"<tr><td>{row['stage']}</td><td>{row['epochs']}</td><td>{row['minutes']}</td><td>{row['pct']}%</td></tr>"
        for row in _stage_summary(predictions)
    )
    metric_rows = "".join(
        f"<tr><td>{_html_escape(k)}</td><td>{_html_escape(v)}</td></tr>"
        for k, v in metrics.items()
    )

    quality_items = ""
    if format_metadata:
        for note in format_metadata.get("conversion_notes", []):
            quality_items += f"<li>{_html_escape(note)}</li>"
        real = format_metadata.get("features_real", [])
        synth = format_metadata.get("features_synthesized", [])
        if real:
            quality_items += f"<li>Real input features: {_html_escape(', '.join(real))}</li>"
        if synth:
            quality_items += f"<li>Estimated/synthesized features: {_html_escape(', '.join(synth))}</li>"

    explanation_section = ""
    if explanation_text:
        explanation_section = f"""
  <section>
    <h2>模型解释</h2>
    <p>{_html_escape(explanation_text)}</p>
  </section>"""

    data_quality_section = ""
    if quality_items:
        data_quality_section = f"""
  <section>
    <h2>数据处理说明</h2>
    <ul>{quality_items}</ul>
  </section>"""

    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>睡眠分期预测报告</title>
  <style>
    body {{ margin: 0; background: #f7f2ea; color: #17211d; font-family: "Microsoft YaHei", Arial, sans-serif; }}
    .wrap {{ max-width: 920px; margin: 0 auto; padding: 40px 22px 72px; }}
    header, section {{ background: #fffaf3; border: 1px solid #eadfce; border-radius: 16px; padding: 24px; margin-bottom: 18px; box-shadow: 0 12px 34px rgba(35, 30, 24, 0.06); }}
    h1 {{ margin: 0 0 8px; font-size: 30px; }}
    h2 {{ margin: 0 0 16px; font-size: 20px; }}
    .meta {{ color: #70695f; font-size: 14px; }}
    table {{ width: 100%; border-collapse: collapse; }}
    th, td {{ padding: 11px 12px; border-bottom: 1px solid #eadfce; text-align: left; }}
    th {{ color: #70695f; font-size: 13px; }}
    img {{ max-width: 100%; border-radius: 12px; }}
    .note {{ color: #5d625d; line-height: 1.8; }}
  </style>
</head>
<body>
<main class="wrap">
  <header>
    <h1>睡眠分期预测报告</h1>
    <div class="meta">生成时间: {now_str} | 数据来源: {_html_escape(upload_filename or "上传数据")}</div>
    <p class="note">本报告只展示模型预测得到的 Wake / NREM / REM 分期结果，以及由分期序列直接汇总得到的基础统计指标。</p>
  </header>

  <section>
    <h2>整晚睡眠分期图</h2>
    <img src="data:image/png;base64,{hypo_b64}" alt="睡眠分期图">
  </section>

  <section>
    <h2>睡眠阶段分布</h2>
    <img src="data:image/png;base64,{dist_b64}" alt="睡眠阶段分布" style="max-width:420px;display:block;margin:auto;">
    <table>
      <thead><tr><th>阶段</th><th>Epoch 数</th><th>时长（分钟）</th><th>占比</th></tr></thead>
      <tbody>{stage_rows}</tbody>
    </table>
  </section>

  <section>
    <h2>由分期结果汇总的基础指标</h2>
    <table>
      <thead><tr><th>指标</th><th>数值</th></tr></thead>
      <tbody>{metric_rows}</tbody>
    </table>
  </section>
{explanation_section}
{data_quality_section}
  <section>
    <h2>重要说明</h2>
    <p class="note">本结果来自可穿戴/睡眠应用数据和机器学习模型推断，不等同于临床多导睡眠图（PSG）诊断。当前模型使用整晚离线数据和上下文窗口进行三分类预测，适合课程展示和算法分析，不作为医疗诊断或治疗建议。</p>
  </section>
</main>
</body>
</html>"""


def generate_docx_report(
    predictions: pd.DataFrame,
    metrics: dict,
    sleep_score: dict = None,
    recommendations: list = None,
    ref_comparison: dict = None,
    shap_importance: pd.DataFrame = None,
    explanation_text: str = "",
    upload_filename: str = "",
    format_metadata: dict = None,
) -> BytesIO:
    """Generate an evidence-focused DOCX report without scores or advice."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return None

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)

    title = doc.add_heading("睡眠分期预测报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"生成时间: {now_str} | 数据来源: {upload_filename or '上传数据'}\n"
        "模型版本: HGB + Context (±2) | 输出: Wake / NREM / REM"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x75, 0x70, 0x68)

    doc.add_paragraph(
        "本报告只展示模型预测得到的睡眠分期结果，以及由分期序列直接汇总得到的基础统计指标。"
    )

    doc.add_heading("整晚睡眠分期图", 1)
    hypo_fig = plot_hypnogram(predictions, figsize=(8, 3))
    hypo_buf = BytesIO()
    hypo_fig.savefig(hypo_buf, format="png", dpi=150, bbox_inches="tight", facecolor=hypo_fig.get_facecolor())
    hypo_buf.seek(0)
    plt.close(hypo_fig)
    doc.add_picture(hypo_buf, width=Inches(6))

    doc.add_heading("睡眠阶段分布", 1)
    dist_fig = plot_stage_distribution(predictions)
    dist_buf = BytesIO()
    dist_fig.savefig(dist_buf, format="png", dpi=150, bbox_inches="tight", facecolor=dist_fig.get_facecolor())
    dist_buf.seek(0)
    plt.close(dist_fig)
    doc.add_picture(dist_buf, width=Inches(3.5))

    stage_table = doc.add_table(rows=1 + 3, cols=4)
    stage_table.style = "Light Shading Accent 1"
    stage_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["阶段", "Epoch 数", "时长（分钟）", "占比"]):
        stage_table.cell(0, j).text = h
    for i, row in enumerate(_stage_summary(predictions), 1):
        stage_table.cell(i, 0).text = row["stage"]
        stage_table.cell(i, 1).text = str(row["epochs"])
        stage_table.cell(i, 2).text = str(row["minutes"])
        stage_table.cell(i, 3).text = f"{row['pct']}%"

    doc.add_heading("由分期结果汇总的基础指标", 1)
    metric_table = doc.add_table(rows=len(metrics) + 1, cols=2)
    metric_table.style = "Light Shading Accent 1"
    metric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metric_table.cell(0, 0).text = "指标"
    metric_table.cell(0, 1).text = "数值"
    for i, (key, value) in enumerate(metrics.items(), 1):
        metric_table.cell(i, 0).text = str(key)
        metric_table.cell(i, 1).text = str(value)

    if format_metadata:
        doc.add_heading("数据处理说明", 1)
        for note in format_metadata.get("conversion_notes", []):
            doc.add_paragraph(str(note), style="List Bullet")
        real = format_metadata.get("features_real", [])
        synth = format_metadata.get("features_synthesized", [])
        if real:
            doc.add_paragraph("Real input features: " + ", ".join(real), style="List Bullet")
        if synth:
            doc.add_paragraph("Estimated/synthesized features: " + ", ".join(synth), style="List Bullet")

    if explanation_text:
        doc.add_heading("模型解释", 1)
        doc.add_paragraph(explanation_text.replace("**", "").replace("*", ""))

    doc.add_heading("重要说明", 1)
    disclaimer = doc.add_paragraph()
    d_run = disclaimer.add_run(
        "本结果来自可穿戴/睡眠应用数据和机器学习模型推断，不等同于临床多导睡眠图（PSG）诊断。"
        "当前模型使用整晚离线数据和上下文窗口进行三分类预测，适合课程展示和算法分析，不作为医疗诊断或治疗建议。"
    )
    d_run.font.size = Pt(9)
    d_run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
