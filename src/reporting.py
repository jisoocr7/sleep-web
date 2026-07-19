"""Evidence-focused HTML and DOCX reports for the safe prototype."""

from __future__ import annotations

import base64
import html
import io
from datetime import datetime, timezone
from functools import lru_cache
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


TEXT = {
    "en": {
        "title": "Wearable Sleep Staging Research Prototype",
        "subtitle": "Evidence-focused prediction report",
        "generated": "Generated (UTC)",
        "model": "Model",
        "model_desc": "HGB context model; 30-s epochs; 9 base features; +/-2 epoch offline context",
        "summary": "Predicted stage summary",
        "stage": "Predicted stage",
        "epochs": "Epochs",
        "minutes": "Duration (min)",
        "percent": "Share of analyzed epochs (%)",
        "hypnogram": "Complete predicted hypnogram",
        "elapsed": "Elapsed time (h)",
        "metrics": "Model-derived exploratory metrics",
        "input_epochs": "Input epochs",
        "analyzed_epochs": "Analyzed epochs",
        "tst": "Predicted total sleep time (min)",
        "sleep_prop": "Predicted sleep proportion within analyzed window (%)",
        "waso": "Predicted wake between first and last sleep epochs (WASO-like, min)",
        "notice": "Derived solely from model-predicted Wake/NREM/REM labels; not independently validated clinical sleep measures.",
        "limits": "Research limitations",
        "limit_items": [
            "Three-class Wake/NREM/REM output; this is not five-stage AASM sleep staging.",
            "Offline inference uses two preceding and two following epochs.",
            "No online training, tuning, calibration, or clinical diagnostic function is performed.",
            "Cross-device generalization is limited; results should not guide clinical decisions.",
        ],
        "footer": "Research use only. Not for diagnosis or clinical decision-making.",
    },
    "zh": {
        "title": "可穿戴睡眠分期研究原型",
        "subtitle": "以预测证据为中心的结果报告",
        "generated": "生成时间（UTC）",
        "model": "模型",
        "model_desc": "HGB 上下文模型；30 秒 epoch；9 个基础特征；前后各 2 个 epoch 的离线上下文",
        "summary": "预测阶段汇总",
        "stage": "预测阶段",
        "epochs": "Epoch 数",
        "minutes": "时长（分钟）",
        "percent": "占全部分析 epoch（%）",
        "hypnogram": "完整预测睡眠阶段图",
        "elapsed": "经过时间（小时）",
        "metrics": "模型派生的探索性指标",
        "input_epochs": "输入 epoch",
        "analyzed_epochs": "分析 epoch",
        "tst": "预测总睡眠时间（分钟）",
        "sleep_prop": "分析窗口内预测睡眠占比（%）",
        "waso": "首末预测睡眠 epoch 之间的清醒时长（WASO-like，分钟）",
        "notice": "以下结果仅由模型预测的 Wake/NREM/REM 标签派生，并非经过独立验证的临床睡眠指标。",
        "limits": "研究限制",
        "limit_items": [
            "输出仅为 Wake/NREM/REM 三分类，不是 AASM 五阶段睡眠分期。",
            "离线推理使用前后各 2 个 epoch 的上下文。",
            "不进行在线训练、调参、校准或临床诊断。",
            "跨设备泛化能力有限，结果不得用于临床决策。",
        ],
        "footer": "仅供研究使用，不用于诊断或临床决策。",
    },
}


def _language(language: str) -> str:
    return "zh" if language == "zh" else "en"


@lru_cache(maxsize=1)
def _cjk_font() -> FontProperties | None:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return FontProperties(fname=str(candidate))
    return None


def _hypnogram_png(stages: list[int], times_seconds: list[float], language: str) -> io.BytesIO:
    language = _language(language)
    text = TEXT[language]
    labels = np.asarray(stages, dtype=int)
    times = np.asarray(times_seconds, dtype=float)
    elapsed_hours = (times - times[0]) / 3600.0 if times.size else np.asarray([], dtype=float)

    fig, axis = plt.subplots(figsize=(10.4, 3.0), dpi=180)
    fig.patch.set_facecolor("white")
    axis.set_facecolor("#fbfaf6")
    axis.step(elapsed_hours, labels, where="post", color="#16372f", linewidth=1.25)
    axis.set_yticks([0, 1, 2], labels=["Wake", "NREM", "REM"])
    axis.invert_yaxis()
    x_label = text["elapsed"]
    label_options = {}
    if language == "zh":
        cjk_font = _cjk_font()
        if cjk_font is not None:
            label_options["fontproperties"] = cjk_font
        else:
            x_label = TEXT["en"]["elapsed"]
    axis.set_xlabel(x_label, fontsize=9, **label_options)
    axis.grid(axis="y", color="#d9ddd8", linewidth=0.7)
    axis.grid(axis="x", color="#ecebe5", linewidth=0.5)
    axis.spines[["top", "right"]].set_visible(False)
    axis.spines[["left", "bottom"]].set_color("#8c9690")
    axis.tick_params(labelsize=8, colors="#34443e")
    axis.margins(x=0)
    fig.tight_layout()

    output = io.BytesIO()
    fig.savefig(output, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    output.seek(0)
    return output


def _stage_rows(payload: dict) -> list[tuple[str, int, float, float]]:
    summary = payload["summary"]["stage_summary"]
    return [
        (stage, summary[stage]["count"], summary[stage]["minutes"], summary[stage]["percent"])
        for stage in ("Wake", "NREM", "REM")
    ]


def generate_html_report(payload: dict, language: str) -> bytes:
    language = _language(language)
    text = TEXT[language]
    chart = _hypnogram_png(payload["stages"], payload["times_seconds"], language)
    chart_b64 = base64.b64encode(chart.getvalue()).decode("ascii")
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    rows = "".join(
        f"<tr><th>{stage}</th><td>{count}</td><td>{minutes:.1f}</td><td>{percent:.1f}</td></tr>"
        for stage, count, minutes, percent in _stage_rows(payload)
    )
    limits = "".join(f"<li>{html.escape(item)}</li>" for item in text["limit_items"])
    summary = payload["summary"]
    metrics = summary["derived_metrics"]

    document = f"""<!doctype html>
<html lang="{language}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(text['subtitle'])}</title>
  <style>
    :root {{ color-scheme: light; --ink:#14221d; --muted:#65736d; --line:#d8ddd8; --paper:#ffffff; --sage:#dce9e2; --coral:#c7775f; }}
    * {{ box-sizing:border-box; }} body {{ margin:0; background:#eef0ec; color:var(--ink); font:15px/1.55 Arial, sans-serif; }}
    main {{ width:min(940px, calc(100% - 32px)); margin:28px auto; background:var(--paper); border:1px solid var(--line); padding:42px; }}
    h1 {{ margin:0; font-size:27px; }} h2 {{ margin:32px 0 12px; font-size:18px; }} p {{ margin:7px 0; }}
    .eyebrow {{ color:#8d4f3d; font-size:12px; font-weight:700; text-transform:uppercase; }} .meta,.notice {{ color:var(--muted); }}
    .model {{ border-left:4px solid #759b87; background:#f2f6f3; padding:12px 16px; margin-top:18px; }}
    table {{ width:100%; border-collapse:collapse; }} th,td {{ border-bottom:1px solid var(--line); padding:10px; text-align:right; }} th:first-child {{ text-align:left; }} thead th {{ background:var(--sage); }}
    img {{ display:block; width:100%; height:auto; border:1px solid var(--line); }}
    .metrics {{ display:grid; grid-template-columns:repeat(3,1fr); gap:10px; }} .metric {{ border:1px solid var(--line); padding:14px; }}
    .metric strong {{ display:block; font-size:21px; }} .notice {{ border:1px solid #e6cbbf; background:#fff7f2; padding:12px; margin-top:14px; }}
    footer {{ margin-top:34px; padding-top:16px; border-top:2px solid var(--ink); font-weight:700; }}
    @media (max-width:700px) {{ main {{ padding:24px; }} .metrics {{ grid-template-columns:1fr; }} }}
  </style>
</head>
<body>
<main>
  <div class="eyebrow">{html.escape(text['subtitle'])}</div>
  <h1>{html.escape(text['title'])}</h1>
  <p class="meta">{html.escape(text['generated'])}: {generated}</p>
  <div class="model"><strong>{html.escape(text['model'])} {html.escape(payload['model_id'])}</strong><br>{html.escape(text['model_desc'])}</div>
  <h2>{html.escape(text['summary'])}</h2>
  <table><thead><tr><th>{html.escape(text['stage'])}</th><th>{html.escape(text['epochs'])}</th><th>{html.escape(text['minutes'])}</th><th>{html.escape(text['percent'])}</th></tr></thead><tbody>{rows}</tbody></table>
  <h2>{html.escape(text['hypnogram'])}</h2>
  <img src="data:image/png;base64,{chart_b64}" alt="{html.escape(text['hypnogram'])}">
  <h2>{html.escape(text['metrics'])}</h2>
  <div class="metrics">
    <div class="metric"><span>{html.escape(text['tst'])}</span><strong>{metrics['predicted_total_sleep_minutes']:.1f}</strong></div>
    <div class="metric"><span>{html.escape(text['sleep_prop'])}</span><strong>{metrics['predicted_sleep_proportion_percent']:.1f}</strong></div>
    <div class="metric"><span>{html.escape(text['waso'])}</span><strong>{metrics['predicted_waso_like_minutes']:.1f}</strong></div>
  </div>
  <p>{html.escape(text['input_epochs'])}: {summary['input_epochs']} | {html.escape(text['analyzed_epochs'])}: {summary['analyzed_epochs']}</p>
  <p class="notice">{html.escape(text['notice'])}</p>
  <h2>{html.escape(text['limits'])}</h2><ul>{limits}</ul>
  <footer>{html.escape(text['footer'])}</footer>
</main>
</body>
</html>"""
    return document.encode("utf-8")


def _set_run_font(run, name="Arial", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor(*color)


def generate_docx_report(payload: dict, language: str) -> io.BytesIO:
    language = _language(language)
    text = TEXT[language]
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    document.core_properties.title = text["subtitle"]
    document.core_properties.subject = "Research-only wearable sleep staging output"
    zoom = document.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(title.add_run(text["title"]), size=19, bold=True, color=(20, 34, 29))
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(subtitle.add_run(text["subtitle"]), size=10, bold=True, color=(141, 79, 61))
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(meta.add_run(f"{text['generated']}: {generated}"), size=9, color=(101, 115, 109))

    model_paragraph = document.add_paragraph()
    _set_run_font(model_paragraph.add_run(f"{text['model']} {payload['model_id']}: "), bold=True)
    _set_run_font(model_paragraph.add_run(text["model_desc"]))

    document.add_heading(text["summary"], level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = [text["stage"], text["epochs"], text["minutes"], text["percent"]]
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        _set_run_font(cell.paragraphs[0].add_run(label), bold=True)
    for stage, count, minutes, percent in _stage_rows(payload):
        cells = table.add_row().cells
        values = [stage, str(count), f"{minutes:.1f}", f"{percent:.1f}"]
        for index, value in enumerate(values):
            _set_run_font(cells[index].paragraphs[0].add_run(value))

    document.add_heading(text["hypnogram"], level=1)
    chart = _hypnogram_png(payload["stages"], payload["times_seconds"], language)
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(chart, width=Inches(6.35))

    document.add_heading(text["metrics"], level=1)
    summary = payload["summary"]
    metrics = summary["derived_metrics"]
    metric_lines = [
        (text["input_epochs"], summary["input_epochs"]),
        (text["analyzed_epochs"], summary["analyzed_epochs"]),
        (text["tst"], f"{metrics['predicted_total_sleep_minutes']:.1f}"),
        (text["sleep_prop"], f"{metrics['predicted_sleep_proportion_percent']:.1f}"),
        (text["waso"], f"{metrics['predicted_waso_like_minutes']:.1f}"),
    ]
    for label, value in metric_lines:
        paragraph = document.add_paragraph()
        _set_run_font(paragraph.add_run(f"{label}: "), bold=True)
        _set_run_font(paragraph.add_run(str(value)))

    notice = document.add_paragraph()
    _set_run_font(notice.add_run(text["notice"]), bold=True, color=(141, 79, 61))

    document.add_heading(text["limits"], level=1)
    for item in text["limit_items"]:
        paragraph = document.add_paragraph(style="List Bullet")
        _set_run_font(paragraph.add_run(item))

    footer = document.add_paragraph()
    footer.paragraph_format.space_before = Pt(12)
    _set_run_font(footer.add_run(text["footer"]), bold=True, color=(20, 34, 29))

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output
